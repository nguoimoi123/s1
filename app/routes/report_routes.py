from flask import Blueprint, request, jsonify, send_file
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

report_bp = Blueprint("report", __name__, url_prefix="/report")


@report_bp.route("/docx", methods=["POST"])
def export_docx():
    data = request.get_json() or {}
    title = data.get("title") or "Meeting Report"
    summary = data.get("summary") or ""
    action_items = data.get("action_items") or []
    key_decisions = data.get("key_decisions") or []
    full_transcript = data.get("full_transcript") or ""

    if not any([summary.strip(), action_items, key_decisions, full_transcript.strip()]):
        return jsonify({"error": "No content to export"}), 400

    doc = Document()
    doc.add_heading(title, level=1)

    if summary.strip():
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(summary)

    if action_items:
        doc.add_heading("Action Items", level=2)
        for item in action_items:
            doc.add_paragraph(str(item), style="List Bullet")

    if key_decisions:
        doc.add_heading("Key Decisions", level=2)
        for item in key_decisions:
            doc.add_paragraph(str(item), style="List Bullet")

    if full_transcript.strip():
        doc.add_heading("Full Transcript", level=2)
        doc.add_paragraph(full_transcript)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="meeting_report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@report_bp.route("/markdown", methods=["POST"])
def export_markdown():
    data = request.get_json() or {}
    title = data.get("title") or "Meeting Report"
    summary = data.get("summary") or ""
    action_items = data.get("action_items") or []
    key_decisions = data.get("key_decisions") or []
    full_transcript = data.get("full_transcript") or ""

    if not any([summary.strip(), action_items, key_decisions, full_transcript.strip()]):
        return jsonify({"error": "No content to export"}), 400

    lines = [f"# {title}"]
    if summary.strip():
        lines += ["## Summary", summary]
    if action_items:
        lines.append("## Action Items")
        lines += [f"- {item}" for item in action_items]
    if key_decisions:
        lines.append("## Key Decisions")
        lines += [f"- {item}" for item in key_decisions]
    if full_transcript.strip():
        lines += ["## Full Transcript", full_transcript]

    content = "\n\n".join(lines)
    buffer = BytesIO(content.encode("utf-8"))
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="meeting_report.md",
        mimetype="text/markdown",
    )


@report_bp.route("/pdf", methods=["POST"])
def export_pdf():
    data = request.get_json() or {}
    title = data.get("title") or "Meeting Report"
    summary = data.get("summary") or ""
    action_items = data.get("action_items") or []
    key_decisions = data.get("key_decisions") or []
    full_transcript = data.get("full_transcript") or ""

    if not any([summary.strip(), action_items, key_decisions, full_transcript.strip()]):
        return jsonify({"error": "No content to export"}), 400

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 50

    def write_line(text, bold=False):
        nonlocal y
        if y < 60:
            pdf.showPage()
            y = height - 50
        pdf.setFont("Helvetica-Bold" if bold else "Helvetica", 12 if bold else 10)
        pdf.drawString(40, y, text[:120])
        y -= 16

    write_line(title, bold=True)

    if summary.strip():
        write_line("Summary", bold=True)
        for line in summary.split("\n"):
            write_line(line)

    if action_items:
        write_line("Action Items", bold=True)
        for item in action_items:
            write_line(f"- {item}")

    if key_decisions:
        write_line("Key Decisions", bold=True)
        for item in key_decisions:
            write_line(f"- {item}")

    if full_transcript.strip():
        write_line("Full Transcript", bold=True)
        for line in full_transcript.split("\n"):
            write_line(line)

    pdf.save()
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="meeting_report.pdf",
        mimetype="application/pdf",
    )
