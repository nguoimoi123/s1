from io import BytesIO
from docx import Document
from ..models.file_model import File
from ..models.folder_model import Folder
from ..models.chunk_model import Chunk
from ..services.plan_service import get_plan_limits, get_user_plan

from openai import OpenAI
import os
from dotenv import load_dotenv

load_dotenv()

client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY")
)

class FileController:

    def slipt_file_to_chunk(text, chunk_size=500, overlap=100):
        chunks = []
        start = 0
        index = 0

        text_length = len(text)

        while start < text_length:
            end = start + chunk_size
            chunk_text = text[start:end]

            chunks.append({
                "chunk_index": index,
                "text": chunk_text.strip()
            })

            index += 1
            start += chunk_size - overlap

        return chunks

    def get_embedding(text: str) -> list[float]:
        response = client.embeddings.create(
            model="text-embedding-3-small",
            input=text
        )
        return response.data[0].embedding

    @staticmethod
    def upload_file(user_id, folder_id, filename, file_type, size, content):
        if not all([user_id, folder_id, filename, file_type, size, content]):
            return {"error": "All file details are required"}, 400

        plan = get_user_plan(user_id)
        limits = get_plan_limits(plan)
        files_limit = limits.get("files_per_folder_limit")
        if files_limit is not None:
            current_count = File.objects(folder_id=folder_id).count()
            if current_count >= files_limit:
                return {
                    "error": "File limit reached for current plan",
                    "plan": plan,
                    "limit": files_limit,
                }, 403

        # Tạo File object
        file = File(
            user_id=user_id,
            folder_id=folder_id,
            filename=filename,
            file_type=file_type,
            size=size,
            content=content
        )
        file.save()

        file_id = str(file.id)

        # CẮT CONTENT THÀNH CHUNK
        chunks = FileController.slipt_file_to_chunk(content)

        # LƯU CHUNK VÀO DB
        chunk_objects = []
        for c in chunks:
            # Tạo embedding
            embedding = FileController.get_embedding(c["text"])
            chunk = Chunk(
                user_id=user_id,
                folder_id=folder_id,
                file_id=file_id,
                chunk_index=c["chunk_index"],
                text=c["text"],
                embedding=embedding
            )
            chunk_objects.append(chunk)

        if chunk_objects:
            Chunk.objects.insert(chunk_objects)

        return {
            "file_id": file_id,
            "filename": filename,
            "total_chunks": len(chunk_objects)
        }, 201

    
    @staticmethod
    def get_files_by_folder(folder_id):
        try:
            folder = Folder.objects.get(id=folder_id)
        except Folder.DoesNotExist:
            return {"error": "Folder not found"}, 404
        files = File.objects(folder_id=folder_id)

        file_list = [{"id": str(file.id), "filename": file.filename, "file_type": file.file_type, "size": file.size, "uploaded_at": file.uploaded_at.isoformat()} for file in files]
        return {
            "folder_name": folder.name,
            "files": file_list}, 200
    
    @staticmethod
    def delete_file(file_id):
        try:
            file = File.objects.get(id=file_id)
        except File.DoesNotExist:
            return {"error": "File not found"}, 404
        
         # Xoá tất cả chunk thuộc file này
        deleted_chunks = Chunk.objects(file_id=str(file.id)).delete()

        # Xoá file
        file.delete()
        return {"message": "File deleted successfully"}, 200

    @staticmethod
    def get_file_for_download(file_id):
        file = File.objects(id=file_id).first()
        if not file:
            return None, None, None, {"error": "File not found"}, 404

        content = file.content or ""
        if not content:
            chunks = Chunk.objects(file_id=str(file.id)).order_by('chunk_index')
            content = "\n".join([c.text for c in chunks])

        if file.file_type.lower() == "docx":
            doc = Document()
            doc.add_heading(file.filename, level=1)
            if content.strip():
                doc.add_paragraph(content)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer, file.filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", None, 200

        text_bytes = content.encode("utf-8")
        buffer = BytesIO(text_bytes)
        buffer.seek(0)
        return buffer, file.filename, "text/plain", None, 200
    